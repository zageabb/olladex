from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

from .workspace import safe_path


ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def inspect_docx(path: Path) -> dict:
    doc = Document(path)
    paragraphs = []
    for index, paragraph in enumerate(doc.paragraphs):
        runs = []
        for run_index, run in enumerate(paragraph.runs):
            color = run.font.color.rgb
            runs.append(
                {
                    "index": run_index,
                    "text": run.text,
                    "bold": bool(run.bold),
                    "italic": bool(run.italic),
                    "underline": bool(run.underline),
                    "font_name": run.font.name or "",
                    "font_size": run.font.size.pt if run.font.size else None,
                    "color": str(color) if color else "",
                }
            )
        paragraphs.append(
            {
                "index": index,
                "text": paragraph.text,
                "style": paragraph.style.name if paragraph.style else "",
                "alignment": _alignment_name(paragraph.alignment),
                "runs": runs,
            }
        )

    tables = [[[cell.text for cell in row.cells] for row in table.rows] for table in doc.tables]
    table_details = []
    for table_index, table in enumerate(doc.tables):
        table_details.append(
            {
                "index": table_index,
                "style": table.style.name if table.style else "",
                "rows": [[cell.text for cell in row.cells] for row in table.rows],
                "row_count": len(table.rows),
                "column_count": len(table.columns),
            }
        )

    sections = []
    for index, section in enumerate(doc.sections):
        sections.append(
            {
                "index": index,
                "orientation": "landscape" if section.orientation == WD_ORIENT.LANDSCAPE else "portrait",
                "page_width_inches": round(section.page_width.inches, 3),
                "page_height_inches": round(section.page_height.inches, 3),
                "top_margin_inches": round(section.top_margin.inches, 3),
                "bottom_margin_inches": round(section.bottom_margin.inches, 3),
                "left_margin_inches": round(section.left_margin.inches, 3),
                "right_margin_inches": round(section.right_margin.inches, 3),
                "header": "\n".join(paragraph.text for paragraph in section.header.paragraphs),
                "footer": "\n".join(paragraph.text for paragraph in section.footer.paragraphs),
            }
        )

    return {
        "kind": "word",
        "paragraphs": paragraphs,
        "headings": [paragraph for paragraph in paragraphs if str(paragraph.get("style", "")).lower().startswith("heading")],
        "tables": tables,
        "table_details": table_details,
        "sections": sections,
        "section_count": len(sections),
        "inline_shapes": [
            {
                "index": index,
                "width_inches": round(shape.width.inches, 3),
                "height_inches": round(shape.height.inches, 3),
            }
            for index, shape in enumerate(doc.inline_shapes)
        ],
    }


def mutate_docx(project: dict, path: Path, operations: list[dict[str, Any]]) -> None:
    doc = Document(path)
    for operation in operations:
        action = operation.get("action")
        if action == "set_paragraph":
            paragraph = _paragraph(doc, operation)
            if "text" in operation:
                paragraph.text = str(operation.get("text", ""))
            if operation.get("style"):
                paragraph.style = str(operation["style"])
            if operation.get("alignment"):
                paragraph.alignment = _alignment(operation["alignment"])
        elif action == "set_paragraph_style":
            _paragraph(doc, operation).style = str(operation.get("style", "Normal"))
        elif action == "set_paragraph_alignment":
            _paragraph(doc, operation).alignment = _alignment(operation.get("alignment"))
        elif action == "set_run":
            paragraph = _paragraph(doc, operation)
            run_index = _index(operation, "run_index", len(paragraph.runs))
            run = paragraph.runs[run_index]
            if "text" in operation:
                run.text = str(operation.get("text", ""))
            if "bold" in operation:
                run.bold = bool(operation.get("bold"))
            if "italic" in operation:
                run.italic = bool(operation.get("italic"))
            if "underline" in operation:
                run.underline = bool(operation.get("underline"))
            if "font_name" in operation:
                run.font.name = str(operation.get("font_name") or "") or None
            if "font_size" in operation:
                size = operation.get("font_size")
                run.font.size = Pt(float(size)) if size not in (None, "") else None
            if "color" in operation:
                color = str(operation.get("color") or "").strip().lstrip("#")
                run.font.color.rgb = RGBColor.from_string(color.upper()) if color else None
        elif action == "append_paragraph":
            paragraph = doc.add_paragraph(str(operation.get("text", "")), style=str(operation.get("style")) if operation.get("style") else None)
            if operation.get("alignment"):
                paragraph.alignment = _alignment(operation["alignment"])
        elif action == "insert_paragraph_after":
            paragraph = _paragraph(doc, operation)
            inserted = _insert_after(paragraph, str(operation.get("text", "")), operation.get("style"))
            if operation.get("alignment"):
                inserted.alignment = _alignment(operation["alignment"])
        elif action == "delete_paragraph":
            paragraph = _paragraph(doc, operation)
            element = paragraph._element
            element.getparent().remove(element)
        elif action == "add_heading":
            level = int(operation.get("level", 1))
            if level < 0 or level > 9:
                raise ValueError("Heading level must be between 0 and 9")
            doc.add_heading(str(operation.get("text", "")), level=level)
        elif action == "set_table_cell":
            table = _table(doc, operation)
            row_index = _index(operation, "row_index", len(table.rows))
            column_index = _index(operation, "column_index", len(table.rows[row_index].cells))
            table.rows[row_index].cells[column_index].text = str(operation.get("text", ""))
        elif action == "add_table":
            rows = max(1, int(operation.get("rows", 2)))
            columns = max(1, int(operation.get("columns", 2)))
            table = doc.add_table(rows=rows, cols=columns)
            if operation.get("style"):
                table.style = str(operation["style"])
            data = operation.get("data")
            if isinstance(data, list):
                for row_index, values in enumerate(data[:rows]):
                    if not isinstance(values, list):
                        continue
                    for column_index, value in enumerate(values[:columns]):
                        table.cell(row_index, column_index).text = str(value if value is not None else "")
        elif action == "add_table_row":
            table = _table(doc, operation)
            row = table.add_row()
            values = operation.get("values")
            if isinstance(values, list):
                for index, value in enumerate(values[: len(row.cells)]):
                    row.cells[index].text = str(value if value is not None else "")
        elif action == "add_table_column":
            table = _table(doc, operation)
            width = float(operation.get("width_inches", 1.25))
            table.add_column(Inches(width))
        elif action == "add_image":
            image_path = str(operation.get("image_path", "")).strip()
            if not image_path:
                raise ValueError("add_image requires image_path")
            image = safe_path(project, image_path)
            if not image.is_file():
                raise ValueError("Image file not found in project")
            width = operation.get("width_inches")
            paragraph_index = operation.get("paragraph_index")
            if paragraph_index is None:
                run = doc.add_paragraph().add_run()
            else:
                paragraph = _paragraph(doc, {"paragraph_index": paragraph_index})
                run = paragraph.add_run()
            run.add_picture(str(image), width=Inches(float(width)) if width else None)
        elif action == "add_hyperlink":
            paragraph = _paragraph(doc, operation)
            text = str(operation.get("text", "")).strip()
            url = str(operation.get("url", "")).strip()
            if not text or not url:
                raise ValueError("add_hyperlink requires text and url")
            _add_hyperlink(paragraph, text, url)
        elif action == "set_section":
            section_index = _index(operation, "section_index", len(doc.sections))
            section = doc.sections[section_index]
            for field in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
                if field in operation:
                    setattr(section, field, Inches(float(operation[field])))
            orientation = operation.get("orientation")
            if orientation:
                next_orientation = str(orientation).lower()
                if next_orientation not in {"portrait", "landscape"}:
                    raise ValueError("orientation must be portrait or landscape")
                landscape = next_orientation == "landscape"
                if landscape != (section.orientation == WD_ORIENT.LANDSCAPE):
                    section.page_width, section.page_height = section.page_height, section.page_width
                section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
        elif action in {"set_header", "set_footer"}:
            section_index = _index(operation, "section_index", len(doc.sections))
            section = doc.sections[section_index]
            container = section.header if action == "set_header" else section.footer
            text = str(operation.get("text", ""))
            if not container.paragraphs:
                container.add_paragraph(text)
            else:
                container.paragraphs[0].text = text
                for paragraph in container.paragraphs[1:]:
                    paragraph.text = ""
        else:
            raise ValueError(f"Unsupported DOCX edit action: {action}")
    doc.save(path)


def _alignment_name(value: Any) -> str:
    if value is None:
        return "left"
    for name, enum_value in ALIGNMENTS.items():
        if value == enum_value:
            return name
    return str(getattr(value, "name", value)).lower()


def _alignment(value: Any):
    key = str(value or "left").strip().lower()
    if key not in ALIGNMENTS:
        raise ValueError("alignment must be left, center, right or justify")
    return ALIGNMENTS[key]


def _paragraph(doc: Document, operation: dict[str, Any]):
    index = _index(operation, "paragraph_index", len(doc.paragraphs))
    return doc.paragraphs[index]


def _table(doc: Document, operation: dict[str, Any]):
    index = _index(operation, "table_index", len(doc.tables))
    return doc.tables[index]


def _index(operation: dict[str, Any], key: str, length: int) -> int:
    try:
        index = int(operation.get(key))
    except (TypeError, ValueError) as exc:
        raise ValueError(f"{key} must be an integer") from exc
    if index < 0 or index >= length:
        raise ValueError(f"{key} is out of range")
    return index


def _insert_after(paragraph, text: str, style: Any = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    if style:
        inserted.style = str(style)
    inserted.add_run(text)
    return inserted


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1768E5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(color)
    properties.append(underline)
    run.append(properties)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
